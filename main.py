from kivy.app  import App
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.splitter import Splitter
from kivy.garden.filebrowser import FileBrowser
from os.path import sep, expanduser, isdir, dirname
from kivy.uix.popup import Popup
import sys
from kivy.properties import ObjectProperty,StringProperty,BooleanProperty
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.dropdown import DropDown
from kivy.uix.scrollview import ScrollView
from kivy.uix.textinput import TextInput
from kivy.uix.colorpicker import ColorPicker
from kivy.core.window import Window
from docx import Document
from docx.shared import Pt
import os
from kivy.graphics import Color, Rectangle,RoundedRectangle

def normal_doc(text):
	document = Document()

	#Institute Name
	p = document.add_paragraph()
	r = p.add_run(text)
	f = r.font
	f.size = Pt(15)

	return document

#Text Input
class MyTextInput(TextInput):
    def __init__(self, **kwargs):
        self.prop_list = {'font_name':"Arial",'font_size':15,'bold':'normal','italic':'normal','color':[0,0,0,1],'la':'down','ca':'normal','ra':'normal','ls':1.0}
        super(MyTextInput, self).__init__(**kwargs)

class PropertiesPanel(BoxLayout):
	line_space = ObjectProperty()
	left_align = ObjectProperty()
	right_align = ObjectProperty()
	center_align = ObjectProperty()
	font_bold = ObjectProperty()
	font_italic = ObjectProperty()
	font_name = ObjectProperty()
	font_size = StringProperty()
	color = ObjectProperty()

	def showlp(self):
		dropdown = CustomDropDown()
		dropdown.bind(on_select=lambda instance, x: setattr(self.line_space, 'text', x))
		dropdown.open(self.line_space)

	def align_text(self):
		if self.left_align.state == 'down':
			App.get_running_app().root.text_input.halign = "left"
			App.get_running_app().root.text_input.prop_list['la'] = 'down'
			App.get_running_app().root.text_input.prop_list['ra'] = 'normal'
			App.get_running_app().root.text_input.prop_list['ca'] = 'normal'
		elif self.right_align.state =='down':
			App.get_running_app().root.text_input.halign = "right"
			App.get_running_app().root.text_input.prop_list['ra'] = 'down'
			App.get_running_app().root.text_input.prop_list['la'] = 'normal'
			App.get_running_app().root.text_input.prop_list['ca'] = 'normal'
		elif self.center_align.state =='down':
			App.get_running_app().root.text_input.halign = 'center'
			App.get_running_app().root.text_input.prop_list['ca'] = 'down'
			App.get_running_app().root.text_input.prop_list['la'] = 'normal'
			App.get_running_app().root.text_input.prop_list['ra'] = 'normal'
		

	def set_line_space(self):
		if self.line_space.text != '':
			App.get_running_app().root.text_input.line_spacing = float(self.line_space.text)
			App.get_running_app().root.text_input.prop_list['ls'] = float(self.line_space.text)

	def set_font_size(self):
		App.get_running_app().root.text_input.font_size = int(self.font_size)
		App.get_running_app().root.text_input.prop_list['font_size'] = int(self.font_size)

	def open_color(self):
		color_box = ColorPicker()
		color_box.color = [0,0,0,1]
		color_box.bind(color=self.set_color)
		self._popup = Popup(title="Text Color", content=color_box,size_hint=(.55, .55))
		self._popup.open()

	def set_color(self,instance,value):
		App.get_running_app().root.text_input.foreground_color  = instance.color
		self.color.background_color = instance.color
		App.get_running_app().root.text_input.prop_list['color'] = instance.color

	def open_fonts(self):
		dropdown = FontDropDown()
		dropdown.bind(on_select=lambda instance, x: setattr(self.font_name, 'text', x))
		dropdown.open(self.font_name)

	def set_font(self):
		App.get_running_app().root.text_input.prop_list['font_name'] = self.font_name.text
		if self.font_bold.state == "down" and self.font_italic.state == "down":
			App.get_running_app().root.text_input.font_name = 'fonts//' + self.font_name.text[0:len(self.font_name.text)].lower().replace(' ','') + "bi.ttf"
			App.get_running_app().root.text_input.prop_list['bold'] = App.get_running_app().root.text_input.prop_list['italic'] = 'down'
		elif self.font_bold.state == "down":
			App.get_running_app().root.text_input.prop_list['bold'] = 'down'
			App.get_running_app().root.text_input.prop_list['italic'] = 'normal'
			App.get_running_app().root.text_input.font_name = 'fonts//' + self.font_name.text[0:len(self.font_name.text)].lower().replace(' ','') + "0b.ttf"
		elif self.font_italic.state == "down":
			App.get_running_app().root.text_input.prop_list['italic'] = 'down'
			App.get_running_app().root.text_input.font_name = 'fonts//' + self.font_name.text[0:len(self.font_name.text)].lower().replace(' ','') + "0i.ttf"
			App.get_running_app().root.text_input.prop_list['bold'] = 'normal'
		else:
			App.get_running_app().root.text_input.font_name = 'fonts//' + self.font_name.text.lower().replace(' ','') + "0r"
			App.get_running_app().root.text_input.prop_list['bold']  = App.get_running_app().root.text_input.prop_list['italic'] = 'normal'

	def reset_panel(self,text_input):
		self.color.foreground_color = text_input.prop_list['color']
		self.font_name.text = text_input.prop_list['font_name']
		self.font_size = str(text_input.prop_list['font_size'])
		self.font_bold.state = text_input.prop_list['bold']
		self.font_italic.state = text_input.prop_list['italic']
		self.line_spacing = text_input.prop_list['ls']
		self.left_align.state = text_input.prop_list['la']
		self.right_align.state = text_input.prop_list['ra']
		self.center_align.state = text_input.prop_list['ca']


class RootWindow(FloatLayout):
	text_input = ObjectProperty()
	path = StringProperty('')
	prop_panel = ObjectProperty()
	appm = ObjectProperty()

	def __init__(self, **kwargs):
		super(RootWindow, self).__init__(**kwargs)

	def open(self,instance):
		path = instance.filename
		if not path == '':
			self._popup.dismiss()
			self.path = path
			if path.split('.')[-1] == 'docx':
				f = open(path, 'rb')
				document = Document(f)
				for	p in document.paragraphs:
					self.text_input.insert_text(p.text)
			else:
				with open(path) as f:
					self.text_input.insert_text(f.read())
		f.close()

	def new(self):
		self.path = ''
		self.text_input.select_all()
		self.text_input.delete_selection()

	def save_as_browser(self):
		if sys.platform == 'win':
			user_path = dirname(expanduser('~')) + sep + 'Documents'
		else:
			user_path = expanduser('~') + sep + 'Documents'
			browser = FileBrowser(select_string='Save',favorites=[(user_path, 'Documents')])
			browser.bind(on_success=self.save_as,on_canceled=self._fbrowser_close)
			self._popup = Popup(title="Save file", content=browser,size_hint=(0.9, 0.9))
			self._popup.open()

	def open_file_browser(self):
		if sys.platform == 'win':
			user_path = dirname(expanduser('~')) + sep + 'Documents'
		else:
			user_path = expanduser('~') + sep + 'Documents'
			browser = FileBrowser(select_string='Open',favorites=[(user_path, 'Documents')])
			browser.bind(on_success=self.open,on_canceled=self._fbrowser_close)
			self._popup = Popup(title="Open file", content=browser,size_hint=(0.9, 0.9))
			self._popup.open()

	def save_as(self, instance):
		path = instance.filename
		if not path == '':
			self._popup.dismiss()
			self.path = path
			self.save()

	def save(self):
		if not self.path == '':
			d = normal_doc(self.text_input.text)
			d.save(self.path.replace('.docx','') + '.docx')
		else:
			self.save_as_browser()

	def open_find_next(self):
		find_box = FindNextBox()
		find_box.bind(find=self.find_next)
		self._popup = Popup(title="Find Next", content=find_box,size_hint=(0.3, 0.25))
		self._popup.open()

	def find_next(self,instance,value):
		start = 0
		end = 0

		current_text = self.text_input.text[instance.cur:]
		start = current_text.find(instance.text_input.text) + instance.cur
		end = start +  len(instance.text_input.text)
		if start >= 0:
			instance.cur = start + len(instance.text_input.text)
			self.text_input.select_text(start, end)

	def open_find(self):
		find_box = FindBox()
		find_box.bind(find=self.find)
		self._popup = Popup(title="Find", content=find_box,size_hint=(0.3, 0.25))
		self._popup.open()

	def find(self,instance,value):
		start = 0
		end = 0
		current_text = self.text_input.text[instance.cur:]
		while(current_text.find(instance.text_input.text)>=0):
			start = current_text.find(instance.text_input.text) + instance.cur
			end = start +  len(instance.text_input.text)
			if start >= 0:
				instance.cur = start + len(instance.text_input.text)
				self.text_input.select_text(start, end)
				self.text_input.selection_color = [1,1,0,.3]
				current_text = self.text_input.text[instance.cur:]

	def open_replace(self):
		replace_box = ReplaceBox()
		replace_box.bind(find1=self.replace)
		replace_box.bind(find2=self.replace_all)
		self._popup = Popup(title="Replace", content=replace_box,size_hint=(0.3, 0.35))
		self._popup.open()

	def replace(self,instance,value):
		start = 0
		end = 0

		current_text = self.text_input.text[instance.cur:]
		start = current_text.find(instance.replace_input.text) + instance.cur
		end = start +  len(instance.replace_input.text)
		if start >= 0:
			instance.cur = start + len(instance.replace_input.text)
			self.text_input.select_text(start, end)
			self.text_input.selection_color = [1,0,0,1]
			self.text_input.copy(data=instance.with_input.text)
			self.text_input.paste()

	def replace_all(self,instance,value):
		start = 0
		end = 0
		current_text = self.text_input.text[instance.cur:]
		while(current_text.find(instance.replace_input.text)>=0):
			start = current_text.find(instance.replace_input.text) + instance.cur
			end = start +  len(instance.replace_input.text)
			if start >= 0:
				instance.cur = start + len(instance.replace_input.text)
				self.text_input.select_text(start, end)
				self.text_input.selection_color = [1,0,0,1]
				self.text_input.copy(data=instance.with_input.text)
				self.text_input.paste()
				current_text = self.text_input.text[instance.cur:]

	def _fbrowser_close(self, instance):
		self._popup.dismiss()

class FindNextBox(BoxLayout):
	text_input = ObjectProperty()
	find = BooleanProperty(False)
	cur = 0

	def changefind(self):
		self.find = not self.find

class FindBox(BoxLayout):
	text_input = ObjectProperty()
	find = BooleanProperty(False)
	cur = 0

	def changefind(self):
		self.find = not self.find

class ReplaceBox(BoxLayout):
	text_input = ObjectProperty()
	find1 = BooleanProperty(False)
	find2 = BooleanProperty(False)
	cur = 0

	def changefind1(self):
		self.find1 = not self.find1
	def changefind2(self):
		self.find2 = not self.find2


class CustomDropDown(DropDown):
    pass

class FontDropDown(DropDown):
    pass

class ScrollList(ScrollView):
	pass

class Editor(App):

	def build(self):
		window = RootWindow()
		return window

	def setTitle(self,instance,value):
		self.title = value

Editor().run()